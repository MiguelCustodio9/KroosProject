from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_background(cell, fill):
    """Define cor de fundo de uma célula"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill)
    cell._element.get_or_add_tcPr().append(shading_elm)

# Criar documento
doc = Document()

# Título principal
title = doc.add_heading('Dicionário de Dados', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitle = doc.add_heading('KroosProject - Sistema de Gestão Desportivo', level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Introdução
doc.add_paragraph('Documento que descreve a estrutura completa da base de dados do KroosProject, incluindo todas as 27 tabelas, suas colunas e relacionamentos.')

doc.add_paragraph()

# Tabelas
tables_info = [
    {
        'num': '1',
        'name': 'UTILIZADOR',
        'desc': 'Armazena informações dos utilizadores ativos do sistema.',
        'columns': [
            ('utilizador_id', 'int(PK)', 'Identificador único do utilizador'),
            ('nome_utilizador', 'varchar(255)', 'Nome de utilizador único para login'),
            ('foto_perfil', 'mediumblob', 'Imagem de perfil do utilizador (blob binário)'),
            ('email_utilizador', 'varchar(255, UNIQUE)', 'Email do utilizador (único no sistema)'),
            ('telefone_utilizador', 'varchar(20, UNIQUE)', 'Número de telefone do utilizador (único)'),
            ('primeiro_nome', 'varchar(50)', 'Primeiro nome/nome próprio'),
            ('último_nome', 'varchar(50)', 'Apelido/último nome'),
            ('data_nascimento', 'date', 'Data de nascimento'),
            ('password', 'varchar(255)', 'Palavra-passe encriptada (MD5 via TRIGGER)'),
            ('tipo_utilizador', 'enum', "Tipo de utilizador: 'admin', 'treinador', 'jogador', 'admin_clube'"),
        ]
    },
    {
        'num': '2',
        'name': 'VALIDAÇÃO_UTILIZADOR',
        'desc': 'Armazena pedidos de registo pendentes de aprovação.',
        'columns': [
            ('validacao_id', 'int(PK)', 'Identificador único do pedido de validação'),
            ('nome_utilizador', 'varchar(255, UNIQUE)', 'Nome de utilizador proposto'),
            ('foto_perfil', 'mediumblob', 'Imagem de perfil proposta'),
            ('email_utilizador', 'varchar(255, UNIQUE)', 'Email proposto para validação'),
            ('telefone_utilizador', 'varchar(20, UNIQUE)', 'Telefone proposto'),
            ('primeiro_nome', 'varchar(50)', 'Primeiro nome proposto'),
            ('último_nome', 'varchar(50)', 'Apelido proposto'),
            ('data_nascimento', 'date', 'Data de nascimento proposta'),
            ('password', 'varchar(255)', 'Palavra-passe em texto (não encriptada ainda)'),
            ('tipo_utilizador', 'enum', "Tipo solicitado: 'admin_clube', 'treinador', 'jogador', '' (vazio)"),
        ]
    },
    {
        'num': '3',
        'name': 'CLUBE',
        'desc': 'Informações gerais dos clubes desportivos.',
        'columns': [
            ('clube_id', 'int(PK)', 'Identificador único do clube'),
            ('nome_clube', 'varchar(100)', 'Nome oficial do clube'),
            ('sigla', 'char(5)', 'Sigla/código do clube (ex: SCP, SLB)'),
            ('logotipo', 'mediumblob', 'Imagem do logótipo do clube'),
            ('cor', 'char(7)', 'Cor principal do clube em formato hexadecimal (#RRGGBB)'),
            ('data_fundação', 'date', 'Data de fundação do clube'),
            ('sede_morada', 'varchar(100)', 'Endereço da sede'),
            ('país_clube', 'varchar(50)', 'País onde o clube está registado'),
            ('cidade_clube', 'varchar(50)', 'Cidade do clube'),
            ('telefone_clube', 'varchar(20)', 'Número de telefone de contacto'),
            ('email_clube', 'varchar(255)', 'Email oficial do clube'),
            ('website_clube', 'varchar(100)', 'Website/URL do clube'),
            ('presidente_clube', 'varchar(100)', 'Nome do presidente do clube'),
            ('instagram_clube', 'varchar(100)', 'Handle do Instagram'),
            ('facebook_clube', 'varchar(100)', 'URL ou nome da página Facebook'),
            ('youtube_clube', 'varchar(100)', 'Canal do YouTube'),
            ('twitter_clube', 'varchar(100)', 'Handle do Twitter'),
            ('tiktok_clube', 'varchar(100)', 'Handle do TikTok'),
            ('código_clube', 'varchar(100, UNIQUE)', 'Código único do clube para identificação'),
        ]
    },
    {
        'num': '4',
        'name': 'ÉPOCA',
        'desc': 'Registos das épocas/temporadas desportivas.',
        'columns': [
            ('epoca_id', 'int(PK)', 'Identificador único da época'),
            ('época', 'varchar(255, UNIQUE)', "Designação da época (ex: '2025/2026')"),
        ]
    },
    {
        'num': '5',
        'name': 'EQUIPA',
        'desc': 'Equipas de um clube em diferentes escalões e épocas.',
        'columns': [
            ('equipa_id', 'int(PK)', 'Identificador único da equipa'),
            ('escalão', 'enum', 'Escalão de competição: S5-S23 ou Seniores'),
            ('hierarquia', 'enum', 'Hierarquia da equipa: A (principal), B, C, D, ... Z'),
            ('epoca_id', 'int(FK)', 'Referência à época desportiva'),
            ('clube_id', 'int(FK)', 'Referência ao clube proprietário'),
        ]
    },
    {
        'num': '6',
        'name': 'ACESSO_EQUIPA',
        'desc': 'Controlo de acesso dos utilizadores às equipas.',
        'columns': [
            ('acesso_id', 'int(PK)', 'Identificador único do acesso'),
            ('equipa_id', 'int(FK)', 'Referência à equipa'),
            ('utilizador_id', 'int(FK)', 'Referência ao utilizador com acesso'),
        ]
    },
    {
        'num': '7',
        'name': 'ESTÁDIO',
        'desc': 'Recintos desportivos do clube.',
        'columns': [
            ('estadio_id', 'int(PK)', 'Identificador único do estádio'),
            ('clube_id', 'int(FK)', 'Referência ao clube proprietário'),
            ('nome_estádio', 'varchar(100, UNIQUE)', 'Nome oficial do estádio'),
            ('capacidade', 'int', 'Número máximo de espectadores'),
        ]
    },
    {
        'num': '8',
        'name': 'JOGADORES',
        'desc': 'Informações pessoais e profissionais dos jogadores.',
        'columns': [
            ('jogador_id', 'int(PK)', 'Identificador único do jogador'),
            ('foto_jogador', 'mediumblob', 'Fotografia do jogador'),
            ('nome_completo', 'varchar(100)', 'Nome completo legal'),
            ('alcunha_jogador', 'varchar(100)', 'Alcunha/apelido desportivo'),
            ('número_favorito', 'enum', 'Número de camisola preferido (1-99)'),
            ('posição_principal', 'enum', 'Posição principal (Guarda-Redes, Defesa Central, etc)'),
            ('posição_secundária', 'enum', 'Posição secundária (pode jogar em mais que uma)'),
            ('data_nascimento', 'date', 'Data de nascimento'),
            ('local_nascimento', 'varchar(100)', 'Localidade de nascimento'),
            ('nacionalidade', 'varchar(100)', 'Nacionalidade'),
            ('país_nascimento', 'varchar(100)', 'País de nascimento'),
            ('pé_preferencial', 'enum', 'Pé dominante: Direito, Esquerdo ou Ambos'),
            ('altura', 'varchar(3)', 'Altura em cm'),
            ('peso', 'varchar(3)', 'Peso em kg'),
            ('instagram', 'varchar(100)', 'Handle do Instagram do jogador'),
            ('facebook', 'varchar(100)', 'Perfil do Facebook'),
            ('twitter', 'varchar(100)', 'Handle do Twitter'),
            ('equipa_id', 'int(FK)', 'Referência à equipa atual'),
        ]
    },
    {
        'num': '9',
        'name': 'ASSIDUIDADE',
        'desc': 'Registos de presença/ausência em treinos.',
        'columns': [
            ('assiduidade_id', 'int(PK)', 'Identificador único do registo'),
            ('treino_id', 'int(FK)', 'Referência ao treino'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador'),
            ('estado', 'enum', "Estado: 'Presente', 'Não Presente Justificado', 'Não Presente Injustificado', 'Lesionado Presente', 'Lesionado Não Presente'"),
            ('justificação_ausência', 'text', 'Motivo da ausência (se aplicável)'),
        ]
    },
    {
        'num': '10',
        'name': 'LESÕES',
        'desc': 'Registos de lesões dos jogadores.',
        'columns': [
            ('lesao_id', 'int(PK)', 'Identificador único da lesão'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador lesionado'),
            ('nome_lesão', 'varchar(100)', 'Nome da lesão (ex: Entorse, Fratura)'),
            ('descrição_lesão', 'varchar(100)', 'Descrição detalhada da lesão'),
            ('tipo_lesão', 'enum', "Classificação: 'Óssea', 'Muscular', 'Ligamentar/Articular', 'Neurológica', 'Cutânea'"),
            ('tempo_recuperação', 'enum', "Tempo estimado: '5 dias - 1 semana' até '+1 ano'"),
            ('estado_lesão', 'enum', "Estado atual: 'Lesionado', 'Em recuperação', 'Em retorno progressivo', 'Recuperado'"),
        ]
    },
    {
        'num': '11',
        'name': 'HISTÓRICO_CARREIRA',
        'desc': 'Histórico de carreira de jogadores por época e clube.',
        'columns': [
            ('carreira_id', 'int(PK)', 'Identificador único do registo'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador'),
            ('epoca_id', 'int(FK)', 'Referência à época'),
            ('clube_id', 'int(FK)', 'Referência ao clube onde jogou'),
            ('jogos', 'int', 'Número de jogos disputados (default: 0)'),
            ('golos_marcados', 'int', 'Total de golos marcados (default: 0)'),
            ('golos_sofridos', 'int', 'Total de golos sofridos (para guarda-redes)'),
            ('assistências', 'int', 'Número de assistências (default: 0)'),
        ]
    },
    {
        'num': '12',
        'name': 'HISTÓRICO_TRANSFERÊNCIA',
        'desc': 'Histórico de transferências de jogadores entre clubes.',
        'columns': [
            ('transferencia_id', 'int(PK)', 'Identificador único da transferência'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador transferido'),
            ('clube_origem_id', 'int(FK)', 'Referência ao clube de origem'),
            ('clube_destino_id', 'int(FK)', 'Referência ao clube de destino'),
            ('valor', 'float', 'Valor da transferência em euros (opcional)'),
        ]
    },
    {
        'num': '13',
        'name': 'VALIDAÇÃO_TRANSFERÊNCIA',
        'desc': 'Transferências pendentes de aprovação.',
        'columns': [
            ('validacao_transferencia_id', 'int(PK)', 'Identificador único'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador em transferência'),
            ('clube_origem_id', 'int(FK)', 'Clube de origem proposto'),
            ('clube_destino_id', 'int(FK)', 'Clube de destino proposto'),
            ('valor', 'float', 'Valor proposto da transferência'),
        ]
    },
    {
        'num': '14',
        'name': 'EXERCÍCIOS',
        'desc': 'Catálogo de exercícios de treino com detalhes técnicos.',
        'columns': [
            ('exercicio_id', 'int(PK)', 'Identificador único do exercício'),
            ('esquema', 'mediumblob', 'Diagrama/esquema visual do exercício'),
            ('estrutura', 'varchar(50)', 'Estrutura do exercício'),
            ('descrição_exercício', 'text', 'Descrição pormenorizada do exercício'),
            ('variantes', 'text', 'Variações/alternativas do exercício'),
            ('fundamentos_ofensivos', 'text', 'Fundamentos ofensivos treinados'),
            ('fundamentos_defensivos', 'text', 'Fundamentos defensivos treinados'),
            ('ações_ofensivas', 'text', 'Ações ofensivas praticadas'),
            ('ações_defensivas', 'text', 'Ações defensivas praticadas'),
            ('duração', 'time', 'Tempo total do exercício'),
            ('repetições', 'int', 'Número de repetições'),
            ('séries', 'int', 'Número de séries'),
            ('pausa_entre_repetições', 'int', 'Pausa entre repetições em segundos'),
            ('pausa_entre_séries', 'int', 'Pausa entre séries em segundos'),
            ('volume_exercício', 'time', 'Volume total (duração × séries)'),
            ('recuperação_para_próximo', 'time', 'Tempo de recuperação antes do próximo exercício'),
        ]
    },
    {
        'num': '15',
        'name': 'PLANO_TREINO',
        'desc': 'Planos de treino com até 10 exercícios.',
        'columns': [
            ('plano_treino_id', 'int(PK)', 'Identificador único do plano'),
            ('exercicio_1_id', 'int(FK)', 'Referência ao exercício 1'),
            ('exercicio_2_id', 'int(FK)', 'Referência ao exercício 2'),
            ('exercicio_3_id', 'int(FK)', 'Referência ao exercício 3'),
            ('exercicio_4_id', 'int(FK)', 'Referência ao exercício 4'),
            ('exercicio_5_id', 'int(FK)', 'Referência ao exercício 5'),
            ('exercicio_6_id', 'int(FK)', 'Referência ao exercício 6'),
            ('exercicio_7_id', 'int(FK)', 'Referência ao exercício 7'),
            ('exercicio_8_id', 'int(FK)', 'Referência ao exercício 8'),
            ('exercicio_9_id', 'int(FK)', 'Referência ao exercício 9'),
            ('exercicio_10_id', 'int(FK)', 'Referência ao exercício 10'),
        ]
    },
    {
        'num': '16',
        'name': 'TREINO',
        'desc': 'Sessões de treino agendadas.',
        'columns': [
            ('treino_id', 'int(PK)', 'Identificador único do treino'),
            ('número_treino', 'int', 'Número sequencial do treino'),
            ('data', 'date', 'Data de realização/agendamento'),
            ('hora', 'time', 'Hora de início'),
            ('conteúdo', 'text', 'Descrição do conteúdo do treino'),
            ('plano_id', 'int(FK)', 'Referência ao plano de treino utilizado'),
            ('observações', 'text', 'Observações adicionais'),
            ('dia_da_semana', 'enum', 'Dia da semana: Segunda a Domingo'),
        ]
    },
    {
        'num': '17',
        'name': 'FASE',
        'desc': 'Fases de competição pré-definidas.',
        'columns': [
            ('fase_id', 'int(PK)', 'Identificador único'),
            ('fase', 'varchar(50)', "Designação da fase (ex: '1ª Jornada', 'Final')"),
        ]
    },
    {
        'num': '18',
        'name': 'COMPETIÇÃO_DEFAULT',
        'desc': 'Modelos de competição reutilizáveis.',
        'columns': [
            ('competicao_default_id', 'int(PK)', 'Identificador único'),
            ('nome_competição', 'varchar(100)', 'Nome do modelo'),
            ('tipo_competição', 'enum', "Tipo: 'Prova a eliminar' ou 'Prova por jornadas'"),
        ]
    },
    {
        'num': '19',
        'name': 'COMPETIÇÃO',
        'desc': 'Competições/torneios da época.',
        'columns': [
            ('competicao_id', 'int(PK)', 'Identificador único'),
            ('nome_competicao_id', 'int(FK)', 'Referência ao modelo de competição padrão'),
            ('época', 'int', 'Época da competição'),
            ('número_fases', 'int', 'Número total de fases'),
            ('vencedor_id', 'int(FK)', 'Referência à equipa vencedora'),
        ]
    },
    {
        'num': '20',
        'name': 'JOGOS',
        'desc': 'Registos de partidas desportivas.',
        'columns': [
            ('jogo_id', 'int(PK)', 'Identificador único do jogo'),
            ('competicao_id', 'int(FK)', 'Referência à competição'),
            ('fase_id', 'int(FK)', 'Referência à fase da competição'),
            ('clube_casa_id', 'int(FK)', 'Clube que joga em casa'),
            ('clube_visitante_id', 'int(FK)', 'Clube visitante'),
            ('data_jogo', 'date', 'Data da partida'),
            ('hora_jogo', 'time', 'Hora de início'),
            ('local_jogo_id', 'int(FK)', 'Referência ao estádio'),
            ('resultado_casa', 'int', 'Golos marcados pelo clube da casa'),
            ('resultado_fora', 'int', 'Golos marcados pelo clube visitante'),
        ]
    },
    {
        'num': '21',
        'name': 'CONVOCATÓRIA',
        'desc': 'Convocações de jogadores para jogos.',
        'columns': [
            ('convocatoria_id', 'int(PK)', 'Identificador único'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador'),
            ('jogo_id', 'int(FK)', 'Referência ao jogo'),
            ('estado', 'enum', "Estado: 'Convocado', 'Não Convocado', 'Lesionado'"),
        ]
    },
    {
        'num': '22',
        'name': 'DETALHES_JOGO',
        'desc': 'Detalhes de participação de jogadores em jogos.',
        'columns': [
            ('detalhes_jogo_id', 'int(PK)', 'Identificador único'),
            ('jogo_id', 'int(FK)', 'Referência ao jogo'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador'),
            ('tipo_detalhes', 'enum', "Situação: '11 Inicial', 'Suplente Não Utilizado', 'Suplente Utilizado'"),
            ('minuto_detalhe', 'enum', 'Minuto de entrada (se aplicável)'),
        ]
    },
    {
        'num': '23',
        'name': 'EVENTOS_JOGO',
        'desc': 'Eventos durante o jogo (golos, cartões, substituições).',
        'columns': [
            ('evento_id', 'int(PK)', 'Identificador único do evento'),
            ('jogo_id', 'int(FK)', 'Referência ao jogo'),
            ('jogador_id', 'int(FK)', 'Jogador envolvido no evento'),
            ('tipo_evento', 'enum', "Tipo: 'Golos', 'Amarelos', 'Vermelhos', 'Substituição'"),
            ('jogador_entrada_id', 'int(FK)', 'Jogador que entra (em substituições)'),
            ('jogador_saida_id', 'int(FK)', 'Jogador que sai (em substituições)'),
            ('jogador_assistencia_id', 'int(FK)', 'Jogador que fez a assistência (em golos)'),
            ('tipo_golo', 'enum', "Tipo de golo: 'Construção Organizada', 'Transição', 'Canto', 'Penálti', etc"),
            ('zona_golo', 'enum', "Zona: 'Dentro da Área' ou 'Fora da Área'"),
            ('zona_corpo_utilizado_golo', 'enum', "Parte do corpo: 'Pé Esquerdo', 'Pé Direito', 'Cabeça', 'Outro'"),
            ('minuto_evento', 'enum', 'Minuto do evento (1-90)'),
        ]
    },
    {
        'num': '24',
        'name': 'ESTATÍSTICAS_JOGO',
        'desc': 'Estatísticas globais de um jogo.',
        'columns': [
            ('stats_id', 'int(PK)', 'Identificador único'),
            ('jogo_id', 'int(FK, UNIQUE)', 'Referência ao jogo (1:1)'),
            ('posse_casa', 'int', 'Percentagem de posse de bola da casa'),
            ('posse_fora', 'int', 'Percentagem de posse de bola da visitante'),
            ('remates_casa', 'int', 'Número total de remates da casa'),
            ('remates_fora', 'int', 'Número total de remates da visitante'),
            ('remates_baliza_casa', 'int', 'Remates enquadrados da casa'),
            ('remates_baliza_fora', 'int', 'Remates enquadrados da visitante'),
            ('grandes_oportunidades_casa', 'int', 'Oportunidades claras de golo da casa'),
            ('grandes_oportunidades_fora', 'int', 'Oportunidades claras de golo da visitante'),
            ('cantos_casa', 'int', 'Número de cantos da casa'),
            ('cantos_fora', 'int', 'Número de cantos da visitante'),
            ('passes_casa', 'int', 'Total de passes da casa'),
            ('passes_fora', 'int', 'Total de passes da visitante'),
            ('passes_certos_casa', 'int', 'Passes bem sucedidos da casa'),
            ('passes_certos_fora', 'int', 'Passes bem sucedidos da visitante'),
        ]
    },
    {
        'num': '25',
        'name': 'ESTATÍSTICAS_JOGADORES',
        'desc': 'Estatísticas individuais de cada jogador em cada jogo.',
        'columns': [
            ('stats_id', 'int(PK)', 'Identificador único'),
            ('jogador_id', 'int(FK)', 'Referência ao jogador'),
            ('jogo_id', 'int(FK)', 'Referência ao jogo'),
            ('minutos_jogados', 'int', 'Tempo de jogo em minutos'),
            ('golos', 'int', 'Golos marcados'),
            ('remates', 'int', 'Total de remates'),
            ('remates_baliza', 'int', 'Remates enquadrados'),
            ('assistências', 'int', 'Assistências'),
            ('passes_chave', 'int', 'Passes que levam a oportunidades'),
            ('passes', 'int', 'Total de passes'),
            ('passes_certos', 'int', 'Passes bem sucedidos'),
            ('cruzamentos', 'int', 'Número de cruzamentos'),
            ('cruzamentos_certos', 'int', 'Cruzamentos bem sucedidos'),
            ('toques_bola', 'int', 'Contactos com a bola'),
            ('dribles', 'int', 'Tentativas de drible'),
            ('dribles_certos', 'int', 'Dribles bem sucedidos'),
            ('perdas', 'int', 'Passes/controlo perdidos'),
            ('desarmes', 'int', 'Tentativas de desarme'),
            ('desarmes_ganhos', 'int', 'Desames bem sucedidos'),
            ('interceções', 'int', 'Bolas interceptadas'),
            ('alívios', 'int', 'Alívios defensivos'),
            ('bloqueios_remate', 'int', 'Bloqueios de remate'),
            ('duelos', 'int', 'Duelos disputados'),
            ('duelos_ganhos', 'int', 'Duelos vencidos'),
            ('faltas_sofridas', 'int', 'Faltas cometidas contra o jogador'),
            ('faltas_feitas', 'int', 'Faltas cometidas pelo jogador'),
            ('amarelos', 'int', 'Cartões amarelos'),
            ('vermelhos', 'int', 'Cartões vermelhos'),
            ('defesas', 'int', 'Defesas (para guarda-redes)'),
            ('remates_baliza_sofridos', 'int', 'Remates enquadrados sofridos (GR)'),
            ('golos_sofridos', 'int', 'Golos sofridos (GR)'),
            ('clean_sheet', 'enum', "Jogo sem sofrer golos: 'Sim' ou 'Não'"),
            ('saídas', 'int', 'Saídas de guarda-redes'),
            ('saídas_eficazes', 'int', 'Saídas bem sucedidas'),
            ('oportunidades_claras_defendidas', 'int', 'Oportunidades claras impedidas'),
            ('class_média', 'decimal', 'Classificação média do desempenho (0-10)'),
        ]
    },
    {
        'num': '26',
        'name': 'EVENTOS_CLUBE',
        'desc': 'Eventos de equipa (treinos, jogos, reuniões, etc).',
        'columns': [
            ('evento_id', 'int(PK)', 'Identificador único'),
            ('equipa_id', 'int(FK)', 'Referência à equipa'),
            ('tipo_evento', 'enum', "Tipo: 'Treino', 'Jogo', 'Reunião Técnico-Tática', 'Sessão de Recuperação', 'Convívio de Equipa', 'Outro'"),
            ('descrição_evento', 'text', 'Descrição detalhada do evento'),
            ('estado_evento', 'enum', "Estado: 'Por realizar', 'Realizado', 'Cancelado', 'Adiado'"),
            ('data_evento', 'date', 'Data do evento'),
        ]
    },
    {
        'num': '27',
        'name': 'MENSAGENS',
        'desc': 'Sistema de mensagens entre utilizadores.',
        'columns': [
            ('mensagem_id', 'int(PK)', 'Identificador único'),
            ('origem_id', 'int(FK)', 'Utilizador que envia'),
            ('destino_id', 'int(FK)', 'Utilizador que recebe'),
            ('conteúdo', 'text', 'Texto da mensagem'),
            ('estado', 'enum', "Estado: 'Lida' ou 'Não Lida'"),
        ]
    },
]

# Adicionar cada tabela
for table_info in tables_info:
    # Heading da tabela
    heading = doc.add_heading(f"{table_info['num']}. {table_info['name']}", level=2)
    
    # Descrição
    desc_para = doc.add_paragraph(table_info['desc'])
    desc_para.style = 'Normal'
    
    # Tabela
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Coluna'
    header_cells[1].text = 'Tipo'
    header_cells[2].text = 'Descrição'
    
    # Formatar header
    for cell in header_cells:
        set_cell_background(cell, 'D9E1F2')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Adicionar linhas
    for col_name, col_type, col_desc in table_info['columns']:
        row_cells = table.add_row().cells
        row_cells[0].text = col_name
        row_cells[1].text = col_type
        row_cells[2].text = col_desc
    
    # Ajustar largura das colunas
    for row in table.rows:
        row.cells[0].width = Inches(1.8)
        row.cells[1].width = Inches(1.5)
        row.cells[2].width = Inches(2.7)
    
    doc.add_paragraph()

# Adicionar resumo final
doc.add_page_break()
summary_heading = doc.add_heading('Resumo Estatístico', level=1)

doc.add_heading('Total de Tabelas: 27', level=3)

doc.add_heading('Tabelas Principais:', level=3)
categories = [
    '• Utilizadores: UTILIZADOR, VALIDACAO_UTILIZADOR (2)',
    '• Estrutura: CLUBE, EPOCA, EQUIPA, ACESSO_EQUIPA, ESTADIO (5)',
    '• Jogadores: JOGADORES, ASSIDUIDADE, LESOES, HISTORICO_CARREIRA, HISTORICO_TRANSFERENCIA, VALIDACAO_TRANSFERENCIA (6)',
    '• Treino: EXERCICIOS, PLANO_TREINO, TREINO (3)',
    '• Competição: FASE, COMPETICAO_DEFAULT, COMPETICAO (3)',
    '• Jogos: JOGOS, CONVOCATORIA, DETALHES_JOGO, EVENTOS_JOGO, ESTATISTICAS_JOGO, ESTATISTICAS_JOGADORES (6)',
    '• Geral: EVENTOS_CLUBE, MENSAGENS (2)',
]

for cat in categories:
    doc.add_paragraph(cat, style='List Bullet')

doc.add_heading('Padrões de Dados', level=2)

doc.add_heading('Chaves Primárias', level=3)
doc.add_paragraph('Todas usam int(11) NOT NULL AUTO_INCREMENT (excepto tabelas de validação e algumas associativas)')
doc.add_paragraph('Formato de nomenclatura: {tabela_singular}_id')

doc.add_heading('Chaves Estrangeiras', level=3)
doc.add_paragraph('Padrão de nomenclatura: fk_id_{tabela_referenciada}')
doc.add_paragraph('Todas com ON UPDATE CASCADE para integridade referencial')

doc.add_heading('Campos ENUM', level=3)
doc.add_paragraph('Estados de entidades (tipo_utilizador, estado_lesão, etc)', style='List Bullet')
doc.add_paragraph('Valores pré-definidos (posições, dias da semana, etc)', style='List Bullet')
doc.add_paragraph('Facilitam validação e economia de espaço', style='List Bullet')

doc.add_heading('Campos BLOB', level=3)
doc.add_paragraph('logotipo (clube), foto_jogador, foto_perfil (utilizador), esquema (exercício)', style='List Bullet')
doc.add_paragraph('Armazenam imagens binárias', style='List Bullet')

# Salvar
output_path = r'c:\xampp\htdocs\KroosProject\Dicionario_de_Dados_KroosProject.docx'
doc.save(output_path)
print(f'✓ Ficheiro Word criado com sucesso: {output_path}')
